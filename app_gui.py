import os
import subprocess
import platform
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from tkinter.scrolledtext import ScrolledText
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from workers.printer import print_file

def remove_table_borders(doc):
    for table in doc.tables:
        tbl = table._element
        tblBorders = tbl.xpath('.//w:tblBorders')
        for tblBorder in tblBorders:
            for border in tblBorder:
                border.attrib.clear()
                border.set(qn('w:val'), 'nil')

def add_cell_padding(doc, padding_value=100):
    for table in doc.tables:
        for cell in table._cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                node = OxmlElement(f'w:{margin}')
                node.set(qn('w:w'), str(padding_value))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

def add_paragraph_spacing(doc, spacing_value=3):
    for para in doc.paragraphs:
        if 'Heading' in para.style.name:
            para.paragraph_format.space_before = Pt(spacing_value / 2)
            para.paragraph_format.space_after = Pt(spacing_value / 2)
        else:
            para.paragraph_format.space_before = Pt(spacing_value)
            para.paragraph_format.space_after = Pt(spacing_value)

def log_message(log_widget, message, success=True):
    symbol = "✔" if success else "✘"
    color = "green" if success else "red"
    log_widget.insert(tk.END, f"{symbol} {message}\n", (color,))
    log_widget.tag_config(color, foreground=color)

def convert_word_to_pdf(word_files, output_dir, print_after_conversion=False, formatted=False, log_widget=None):
    pdf_files = []
    for word_file in word_files:
        try:
            word_file_to_convert = word_file
            if formatted:
                doc = Document(word_file)
                remove_table_borders(doc)
                add_cell_padding(doc)
                add_paragraph_spacing(doc)
                temp_word_file = word_file.replace('.docx', '_formatted.docx')
                doc.save(temp_word_file)
                if log_widget:
                    log_message(log_widget, f"File formattato salvato come: {temp_word_file}")
                word_file_to_convert = temp_word_file

            pdf_file = os.path.join(output_dir, os.path.splitext(os.path.basename(word_file_to_convert))[0] + '.pdf')
            
            if platform.system() == "Windows":
                soffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
                command = [soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, word_file_to_convert]
            else:
                command = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, word_file_to_convert]
                
            if log_widget:
                log_message(log_widget, f"Esecuzione comando: {' '.join(command)}")

            result = subprocess.run(command, capture_output=True)

            if log_widget:
                log_message(log_widget, f"Output comando: {result.stdout.decode()}")
                log_message(log_widget, f"Errore comando: {result.stderr.decode()}")

            if result.returncode != 0:
                raise RuntimeError(f"Errore nella conversione di {word_file} in PDF: {result.stderr.decode()}")

            if os.path.exists(pdf_file):
                pdf_files.append(pdf_file)
                if log_widget:
                    log_message(log_widget, f"Convertito {word_file} in {pdf_file}", success=True)
            else:
                raise RuntimeError(f"File convertito {pdf_file} non trovato")

        except Exception as e:
            if log_widget:
                log_message(log_widget, f"Conversione fallita: {word_file}\nErrore: {str(e)}", success=False)
            else:
                print(f"✘ Conversione fallita: {word_file}\nErrore: {str(e)}\n")

    if print_after_conversion:
        for pdf_file in pdf_files:
            print_file(pdf_file)

def select_files():
    files = filedialog.askopenfilenames(filetypes=[("File Word", "*.docx")])
    if files:
        file_list.extend(files)
        update_file_list()

def select_folder():
    folder = filedialog.askdirectory()
    if folder:
        for root, _, files in os.walk(folder):
            for file in files:
                if file.endswith(".docx"):
                    file_list.append(os.path.join(root, file))
        update_file_list()

def update_file_list():
    listbox.delete(0, tk.END)
    for f in file_list:
        listbox.insert(tk.END, f)

def clear_file_list():
    file_list.clear()
    update_file_list()

def select_output_folder():
    global output_directory, output_folder_name
    output_directory = filedialog.askdirectory()
    if output_directory:
        output_folder_name = simpledialog.askstring("Nome cartella di output", "Inserisci il nome della cartella di output (default: convertyPDF):", initialvalue="convertyPDF")
        if not output_folder_name:
            output_folder_name = "convertyPDF"
        output_directory = os.path.join(output_directory, output_folder_name)
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
        output_dir_label.config(text=f"Directory di output: {output_directory}")

def convert_files(print_after=False, formatted=False):
    if not file_list:
        messagebox.showwarning("Nessun file selezionato", "Seleziona uno o più file Word da convertire.")
        return
    if not output_directory:
        messagebox.showwarning("Nessuna directory di output selezionata", "Seleziona una directory di output.")
        return

    log_widget.delete(1.0, tk.END)  # Clear previous log
    convert_word_to_pdf(file_list, output_directory, print_after, formatted, log_widget)

app = tk.Tk()
app.title("ConvertifyPDF")
app.geometry("700x500")

file_list = []
output_directory = ""
output_folder_name = "convertyPDF"

frame = tk.Frame(app)
frame.pack(pady=10)

listbox = tk.Listbox(frame, width=80, height=10)
listbox.pack(side=tk.LEFT, padx=(0, 10))

scrollbar = tk.Scrollbar(frame)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

listbox.config(yscrollcommand=scrollbar.set)
scrollbar.config(command=listbox.yview)

button_frame = tk.Frame(app)
button_frame.pack(pady=10)

select_files_button = tk.Button(button_frame, text="Seleziona file Word", command=select_files)
select_files_button.grid(row=0, column=0, padx=5)

select_folder_button = tk.Button(button_frame, text="Seleziona cartella", command=select_folder)
select_folder_button.grid(row=0, column=1, padx=5)

clear_button = tk.Button(button_frame, text="Pulisci lista", command=clear_file_list)
clear_button.grid(row=0, column=2, padx=5)

output_dir_button = tk.Button(app, text="Seleziona directory di output", command=select_output_folder)
output_dir_button.pack(pady=5)

output_dir_label = tk.Label(app, text="Directory di output: Non selezionata")
output_dir_label.pack(pady=5)

convert_button = tk.Button(app, text="Converti", command=lambda: convert_files(print_after=False, formatted=False))
convert_button.pack(pady=5)

convert_print_button = tk.Button(app, text="Converti e stampa", command=lambda: convert_files(print_after=True, formatted=False))
convert_print_button.pack(pady=5)

convert_format_button = tk.Button(app, text="Converti con formattazione", command=lambda: convert_files(print_after=False, formatted=True))
convert_format_button.pack(pady=5)

convert_format_print_button = tk.Button(app, text="Converti con formattazione e stampa", command=lambda: convert_files(print_after=True, formatted=True))
convert_format_print_button.pack(pady=5)

log_widget = ScrolledText(app, width=80, height=10, wrap=tk.WORD)
log_widget.pack(pady=10)

app.mainloop()
