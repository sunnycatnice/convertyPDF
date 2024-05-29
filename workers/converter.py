import subprocess
import os
import platform
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from .printer import print_file

def remove_table_borders(doc):
    for table in doc.tables:
        tbl = table._element
        tblBorders = tbl.xpath('.//w:tblBorders')
        for tblBorder in tblBorders:
            for border in tblBorder:
                border.attrib.clear()
                border.set(qn('w:val'), 'nil')

def add_cell_padding(doc, padding_value=100):
    for table in doc.tables:
        for cell in table._cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                node = OxmlElement(f'w:{margin}')
                node.set(qn('w:w'), str(padding_value))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

def add_paragraph_spacing(doc, spacing_value=3):
    for para in doc.paragraphs:
        if 'Heading' in para.style.name:
            para.paragraph_format.space_before = Pt(spacing_value / 2)
            para.paragraph_format.space_after = Pt(spacing_value / 2)
        else:
            para.paragraph_format.space_before = Pt(spacing_value)
            para.paragraph_format.space_after = Pt(spacing_value)

def convert_word_to_pdf(word_files, print_after_conversion=False, formatted=False, log_widget=None):
    pdf_files = []
    for word_file in word_files:
        try:
            if formatted:
                doc = Document(word_file)
                remove_table_borders(doc)
                add_cell_padding(doc)
                add_paragraph_spacing(doc)
                temp_word_file = word_file.replace('.docx', '_formatted.docx')
                doc.save(temp_word_file)
                if log_widget:
                    log_widget.insert(tk.END, f"Formatted file saved as: {temp_word_file}\n")
                word_file_to_convert = temp_word_file
            else:
                word_file_to_convert = word_file

            pdf_file = word_file_to_convert.replace('.docx', '.pdf')
            
            if platform.system() == "Windows":
                soffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
                command = [soffice_path, '--headless', '--convert-to', 'pdf', word_file_to_convert]
            else:
                command = ['soffice', '--headless', '--convert-to', 'pdf', word_file_to_convert]
                
            if log_widget:
                log_widget.insert(tk.END, f"Running command: {' '.join(command)}\n")

            result = subprocess.run(command, capture_output=True)

            if log_widget:
                log_widget.insert(tk.END, f"Command output: {result.stdout.decode()}\n")
                log_widget.insert(tk.END, f"Command error: {result.stderr.decode()}\n")
                
            if result.returncode != 0:
                raise RuntimeError(f"Error converting {word_file} to PDF: {result.stderr.decode()}")

            converted_file = word_file_to_convert.replace('.docx', '.pdf')
            if os.path.exists(converted_file):
                os.rename(converted_file, pdf_file)
                pdf_files.append(pdf_file)
                if log_widget:
                    log_widget.insert(tk.END, f"Converted {word_file} to {pdf_file}\n")
            else:
                raise RuntimeError(f"Converted file {converted_file} not found")

        except Exception as e:
            if log_widget:
                log_widget.insert(tk.END, f"✘ Conversion failed: {word_file}\nError: {str(e)}\n")
            else:
                print(f"✘ Conversion failed: {word_file}\nError: {str(e)}\n")

    if print_after_conversion:
        for pdf_file in pdf_files:
            print_file(pdf_file)
